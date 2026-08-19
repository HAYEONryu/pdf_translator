"""HTML(대조본) / DOCX(번역본만) 내보내기 (SPEC.md §7.3). 폰트는 맑은 고딕 기본."""
import base64
import html as html_lib
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

from app.extract.render import render_crop_png, render_page_png
from app.storage import source_pdf_path

_TRANSLATABLE = ("heading", "paragraph", "list", "table")


def export_html(doc_sha: str, doc_title: str, pages: list[dict]) -> str:
    """원문 이미지 + 번역 텍스트를 나란히 보여주는 정적 HTML (클릭 대조 없음, 단순 열람용)."""
    pdf_path = str(source_pdf_path(doc_sha))
    parts = [
        "<!doctype html><html lang='ko'><head><meta charset='utf-8'>",
        f"<title>{html_lib.escape(doc_title)} - 대조본</title>",
        "<style>"
        "body{font-family:'Malgun Gothic',sans-serif;margin:0;background:#eee}"
        ".page{display:flex;gap:16px;padding:16px;border-bottom:8px solid #ccc;background:#fff}"
        ".src,.tgt{flex:1;min-width:0}"
        ".src img{max-width:100%;display:block}"
        "table{border-collapse:collapse;width:100%;font-size:13px}"
        "td{border:1px solid #999;padding:4px 8px}"
        ".page-label{color:#888;font-size:12px;margin-top:4px}"
        ".formula-crop{display:block;max-width:100%;background:#f6f6f6;border:1px solid #e0e0e0;"
        "border-radius:3px;padding:4px}"
        "</style></head><body>",
    ]
    for pg in pages:
        img_b64 = base64.b64encode(render_page_png(pdf_path, pg["page_no"])).decode()
        parts.append(f"<div class='page'><div class='src'><img src='data:image/png;base64,{img_b64}'>")
        parts.append(f"<div class='page-label'>p.{pg['page_no']:03d}</div></div><div class='tgt'>")
        for b in pg["blocks"]:
            if b["type"] == "figure":
                # PUA 폰트로 깨지는 텍스트 대신 원본 그대로 보이는 이미지로 잘라 넣는다.
                crop_b64 = base64.b64encode(render_crop_png(pdf_path, pg["page_no"], b["bbox"])).decode()
                parts.append(f"<img class='formula-crop' src='data:image/png;base64,{crop_b64}'>")
                continue
            if b["type"] not in _TRANSLATABLE:
                continue
            if b["type"] == "table" and b["table"] and b["table"]["cells_ko"]:
                parts.append("<table>")
                for row in b["table"]["cells_ko"]:
                    parts.append("<tr>" + "".join(f"<td>{html_lib.escape(c or '')}</td>" for c in row) + "</tr>")
                parts.append("</table>")
            elif b["ko"]:
                tag = "h3" if b["type"] == "heading" else "p"
                parts.append(f"<{tag}>{html_lib.escape(b['ko'])}</{tag}>")
        parts.append("</div></div>")
    parts.append("</body></html>")
    return "".join(parts)


def export_docx(doc_sha: str, doc_title: str, pages: list[dict]) -> bytes:
    """번역본만. 표는 python-docx 표로, 나머지는 문단으로."""
    pdf_path = str(source_pdf_path(doc_sha))
    document = Document()
    document.styles["Normal"].font.name = "맑은 고딕"
    document.styles["Normal"].font.size = Pt(11)

    document.add_heading(doc_title, level=1)
    for pg in pages:
        document.add_heading(f"p.{pg['page_no']:03d}", level=2)
        for b in pg["blocks"]:
            if b["type"] == "figure":
                # PUA 폰트로 깨지는 텍스트 대신 원본 그대로 보이는 이미지로 잘라 넣는다.
                crop = render_crop_png(pdf_path, pg["page_no"], b["bbox"])
                # 폭을 6in으로 고정하면 원래 좁고 납작한 크롭(예: 라벨 한 줄)이 몇 배로
                # 확대돼 흐릿하게 보인다 (실사용 중 발견) — bbox 실제 폭(pt→in)을 쓰고
                # 본문 폭(6in)보다 클 때만 줄인다. 확대는 절대 하지 않는다.
                natural_width_in = (b["bbox"][2] - b["bbox"][0]) / 72.0
                document.add_picture(BytesIO(crop), width=Inches(min(natural_width_in, 6.0)))
                continue
            if b["type"] not in _TRANSLATABLE:
                continue
            if b["type"] == "table" and b["table"] and b["table"]["cells_ko"]:
                cells = b["table"]["cells_ko"]
                table = document.add_table(rows=len(cells), cols=len(cells[0]) if cells else 0)
                table.style = "Table Grid"
                for r, row in enumerate(cells):
                    for c, val in enumerate(row):
                        table.cell(r, c).text = val or ""
            elif b["ko"]:
                document.add_paragraph(b["ko"], style="Heading 3" if b["type"] == "heading" else "Normal")

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()
